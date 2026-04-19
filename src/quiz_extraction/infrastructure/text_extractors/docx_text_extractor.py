import io

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.quiz_extraction.application.ports.text_extractor_port import TextExtractorPort


def _iter_block_items(document):
    """Yield Paragraph and Table elements in document order.

    python-docx exposes `document.paragraphs` and `document.tables` as
    separate lists, losing the original order. Walking the underlying
    OOXML tree preserves sequence, which matters for quiz extraction
    where a matching question may be followed by its answer table.
    """
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


class DocxTextExtractor(TextExtractorPort):
    def extract(self, file_bytes: bytes, filename: str) -> str:
        document = Document(io.BytesIO(file_bytes))
        lines: list[str] = []

        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    lines.append(text)
            elif isinstance(block, Table):
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))

        return "\n".join(lines).strip()
